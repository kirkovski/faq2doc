# BELOW ARE REQUIRED LIBRARIES, install at terminal
# pip install pandas
# pip install openpyxl
# pip install python-docx
import pandas as pd
from docx import Document

# Read Excel data 
df = pd.read_excel('input.xlsx')  # read excel into DataFrame

# Create a new Word document 
doc = Document()

# For each row in the DataFrame
for index, row in df.iterrows():

    # Add a paragraph for each question and answer
    para = doc.add_paragraph()

    # Add the question
    run = para.add_run(f"Question {row['S. No.']}: {row['Question']}")
    run.bold = True

    # Add the answer
    doc.add_paragraph(f"Answer: {row['Answer']}")

    # Add dash marks to separate each question
    doc.add_paragraph("---")

# Save the Word document
doc.save('output.docx')
